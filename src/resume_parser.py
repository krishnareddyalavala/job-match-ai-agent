# src/resume_parser.py
"""
Handles reading and parsing resume from .docx or .pdf files.
"""

import os
from pathlib import Path
from docx import Document
import pypdf


def extract_text_from_docx(file_path: str) -> str:
    """Extract full text from a .docx file."""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    # Also extract from tables (skills table etc.)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text.strip())
    return "\n".join(full_text)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract full text from a .pdf file."""
    reader = pypdf.PdfReader(file_path)
    full_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            full_text.append(text)
    return "\n".join(full_text)


def parse_resume(file_path: str) -> str:
    """
    Parse resume from file. Supports .docx and .pdf.
    Returns the full text content.
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Please upload .docx or .pdf")


def chunk_resume(resume_text: str, chunk_size: int = 500) -> list[str]:
    """
    Split resume text into chunks for embedding/RAG.
    Simple paragraph-based chunking.
    """
    paragraphs = [p.strip() for p in resume_text.split("\n") if p.strip()]
    chunks = []
    current_chunk = []
    current_length = 0

    for para in paragraphs:
        para_len = len(para.split())
        if current_length + para_len > chunk_size and current_chunk:
            chunks.append("\n".join(current_chunk))
            current_chunk = [para]
            current_length = para_len
        else:
            current_chunk.append(para)
            current_length += para_len

    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks
