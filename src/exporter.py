# src/exporter.py
"""
Exports the rewritten resume content to a .docx file.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_docx(
    original_resume_path: str,
    rewritten_summary: str,
    rewritten_bullets: list[str],
    job_title: str,
    output_dir: str = "outputs"
) -> str:
    """
    Creates a new .docx with the rewritten summary and bullets
    appended as a clearly labeled 'JD-Aligned Version' section.
    """
    os.makedirs(output_dir, exist_ok=True)

    # Start from original resume if provided, else blank doc
    if original_resume_path and os.path.exists(original_resume_path):
        doc = Document(original_resume_path)
    else:
        doc = Document()

    # Add a page break before the new section
    doc.add_page_break()

    # Section header
    header = doc.add_heading(f"✦ JD-Aligned Version — {job_title}", level=1)
    header.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph(
        f"Generated on: {datetime.now().strftime('%B %d, %Y %H:%M')}"
    ).italic = True

    # Rewritten Summary
    doc.add_heading("Professional Summary (Rewritten)", level=2)
    p = doc.add_paragraph(rewritten_summary)
    p.paragraph_format.space_after = Pt(12)

    # Rewritten Bullets
    doc.add_heading("Key Experience Bullets (JD-Aligned)", level=2)
    for bullet in rewritten_bullets:
        bp = doc.add_paragraph()
        bp.add_run(f"• {bullet}")

    # Save
    safe_title = job_title.replace(" ", "_").replace("/", "-")[:40]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"resume_aligned_{safe_title}_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    return output_path


def export_report_to_txt(report_text: str, output_dir: str = "outputs") -> str:
    """Save the full text report to a .txt file."""
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"match_report_{timestamp}.txt"
    output_path = os.path.join(output_dir, filename)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(report_text)
    return output_path
